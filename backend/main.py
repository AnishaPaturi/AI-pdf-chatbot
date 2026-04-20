import os
import shutil
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Form, Body
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from typing import List

# Database & Auth
from database import init_db, get_db, engine
from models import Base, Document, QueryLog, Highlight, Note
from auth import (
    UserRegister, UserLogin, TokenResponse, UserResponse,
    create_access_token, decode_token, register_user, 
    authenticate_user, get_user_by_id
)

# LangChain & RAG
from langchain_openai import ChatOpenAI 
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel
from dotenv import load_dotenv

# Load environment variables
load_dotenv(override=True)

# --- FastAPI App Setup ---
app = FastAPI(
    title="AI PDF Chatbot API",
    description="Backend API with Auth, Multi-PDF, and Strict LLM Scoping",
    version="2.0"
)

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Configuration ---
UPLOAD_DIR = "./uploads"
CHROMA_DB_DIR = "./chroma_db"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Initialize database on startup
@app.on_event("startup")
async def startup_event():
    """Initialize database tables on startup."""
    init_db()
    print("✓ Database initialized")

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# Vector store cache
vector_store_cache = {}

def get_or_create_vector_store(user_id: int):
    """Get or create a vector store for a specific user."""
    user_chroma_dir = os.path.join(CHROMA_DB_DIR, f"user_{user_id}")
    
    if user_id not in vector_store_cache:
        if os.path.exists(user_chroma_dir):
            vector_store_cache[user_id] = Chroma(
                persist_directory=user_chroma_dir,
                embedding_function=embeddings
            )
        else:
            vector_store_cache[user_id] = None
    
    return vector_store_cache[user_id]

# --- Security ---
security = HTTPBearer()

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security), db: Session = Depends(get_db)):
    """Extract and validate JWT token from request headers."""
    token = credentials.credentials
    payload = decode_token(token)
    
    if payload is None:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    user = get_user_by_id(db, payload["user_id"])
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    
    return user

# --- Pydantic Models ---
class ChatRequest(BaseModel):
    query: str
    document_ids: List[int] = None  # Optional: filter by specific documents

class ChatResponse(BaseModel):
    answer: str

class SummaryResponse(BaseModel):
    summary: str

# --- Auth Endpoints ---
@app.post("/register", response_model=TokenResponse)
async def register(user_data: UserRegister, db: Session = Depends(get_db)):
    """Register a new user and return access token."""
    new_user = register_user(db, user_data)
    
    if new_user is None:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    token = create_access_token(new_user.id, new_user.email)
    return TokenResponse(
        access_token=token,
        user_id=new_user.id,
        email=new_user.email,
        name=new_user.name
    )

@app.post("/login", response_model=TokenResponse)
async def login(credentials: UserLogin, db: Session = Depends(get_db)):
    """Authenticate user and return access token."""
    user = authenticate_user(db, credentials.email, credentials.password)
    
    if user is None:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    token = create_access_token(user.id, user.email)
    return TokenResponse(
        access_token=token,
        user_id=user.id,
        email=user.email,
        name=user.name
    )

@app.get("/me", response_model=UserResponse)
async def get_current_user_info(
    current_user = Depends(get_current_user)
):
    """Get current logged-in user info."""
    return UserResponse.from_orm(current_user)

# --- PDF Upload Endpoints ---
@app.post("/upload")
async def upload_pdfs(
    files: List[UploadFile] = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload multiple PDF files for a specific user.
    Each file is vectorized and stored in a user-specific vector store.
    """
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="No files provided")
    
    user_id = current_user.id
    user_upload_dir = os.path.join(UPLOAD_DIR, f"user_{user_id}")
    os.makedirs(user_upload_dir, exist_ok=True)
    
    user_chroma_dir = os.path.join(CHROMA_DB_DIR, f"user_{user_id}")
    os.makedirs(user_chroma_dir, exist_ok=True)
    
    results = []
    
    try:
        for file in files:
            # Validate file is PDF
            if not file.filename.endswith('.pdf'):
                continue
            
            # Save file
            file_path = os.path.join(user_upload_dir, file.filename)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            print(f"✓ File uploaded: {file.filename} for user {user_id}")
            
            # Load and process PDF
            try:
                loader = PyPDFLoader(file_path)
                documents = loader.load()
                
                # Split into chunks
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200
                )
                chunks = text_splitter.split_documents(documents)
                
                # Store in user's vector database - merge with existing if present
                existing_store = get_or_create_vector_store(user_id)
                if existing_store is not None:
                    # Add to existing vector store
                    existing_store.add_documents(chunks)
                    vector_store = existing_store
                else:
                    # Create new vector store
                    vector_store = Chroma.from_documents(
                        documents=chunks,
                        embedding=embeddings,
                        persist_directory=user_chroma_dir
                    )
                
                # Update cache
                vector_store_cache[user_id] = vector_store
                
                # Save document metadata to database
                doc_record = Document(
                    user_id=user_id,
                    filename=file.filename,
                    original_filename=file.filename,
                    file_path=file_path,
                    chunk_count=len(chunks)
                )
                db.add(doc_record)
                db.commit()
                db.refresh(doc_record)
                
                results.append({
                    "filename": file.filename,
                    "chunks": len(chunks),
                    "document_id": doc_record.id,
                    "status": "success"
                })
                
            except Exception as e:
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": str(e)
                })
                print(f"✗ Error processing {file.filename}: {e}")
        
        return {
            "user_id": user_id,
            "files_processed": len([r for r in results if r["status"] == "success"]),
            "results": results
        }
        
    except Exception as e:
        print(f"Error during upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Chat Endpoint ---
@app.post("/chat", response_model=ChatResponse)
async def chat_with_pdf(
    request: ChatRequest,
    current_user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Chat with uploaded PDFs using strict RAG context scoping.
    The AI ONLY answers based on the uploaded documents.
    """
    query = request.query
    user_id = current_user.id

    # Get user's vector store
    vector_store = get_or_create_vector_store(user_id)

    if vector_store is None:
        raise HTTPException(
            status_code=400,
            detail="No documents uploaded yet. Please upload PDFs first."
        )

    openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
    if not openrouter_api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENROUTER_API_KEY not configured"
        )

    try:
        # Initialize LLM - use a text-only model
        llm = ChatOpenAI(
            api_key=openrouter_api_key,
            base_url="https://openrouter.ai/api/v1",
            model="google/gemini-2.0-flash-001",
            temperature=0.1,
            max_tokens=4096,
        )

        # STRICT SYSTEM PROMPT - Prevents hallucination
        system_prompt = (
            "You are a strict data-extraction AI. You MUST ONLY answer using the provided Context. "
            "If the answer or topic is not explicitly found in the Context provided below, you MUST reply with: "
            "'I can only answer questions based on the uploaded documents.' "
            "DO NOT use your internal knowledge under any circumstance. "
            "DO NOT make up information. ONLY use what is in the context.\n\n"
            "Context:\n"
            "{context}"
        )

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}"),
        ])

        # Create RAG chain
        # Detect summary queries for better coverage
        query_lower = query.lower()
        k_docs = 20 if (
            "summary" in query_lower
            or "summarize" in query_lower
            or "tl" in query_lower
        ) else 20  # Increased for quoted text search

        # Extract quoted text for better retrieval
        search_query = query
        if "regarding this text:" in query_lower:
            # Extract the quoted text for better search
            import re
            match = re.search(r'regarding this text:?"([^"]+)"', query, re.IGNORECASE)
            if match:
                quoted_text = match.group(1).strip()
                # Use quoted text as the search query for better context
                search_query = quoted_text
                print(f"Using quoted text for retrieval: {quoted_text[:50]}...")

        retriever = vector_store.as_retriever(search_kwargs={"k": k_docs})

        def format_docs(docs):
            return "\n\n".join(doc.page_content for doc in docs)

        rag_chain = (
            {"context": retriever | format_docs, "input": RunnablePassthrough()}
            | prompt
            | llm
            | StrOutputParser()
        )

        # Generate answer
        print(f"User {user_id} querying: {search_query[:50]}...")
        final_answer = rag_chain.invoke(search_query)

        # Log query to database
        document_ids = request.document_ids

        if not document_ids:
            # Use first document if not specified
            first_doc = db.query(Document).filter(
                Document.user_id == user_id
            ).first()

            document_ids = [first_doc.id] if first_doc else [None]

        for doc_id in document_ids:
            if doc_id:
                query_log = QueryLog(
                    user_id=user_id,
                    document_id=doc_id,
                    query=query,
                    response=final_answer
                )
                db.add(query_log)

        db.commit()

        return ChatResponse(answer=final_answer)

    except Exception as e:
        print(f"Error during chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
    
# --- PDF File Serving ---
@app.get("/document/{document_id}/file")
async def get_document_file(
    document_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get the PDF file content for viewing."""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    from fastapi.responses import FileResponse
    return FileResponse(
        document.file_path,
        media_type='application/pdf',
        filename=document.original_filename
    )

# --- Summary Endpoint ---
@app.post("/summary", response_model=SummaryResponse)
async def get_document_summary(
    document_id: int = None,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Generate a summary of the uploaded PDF(s) using strict RAG context scoping.
    """
    user_id = current_user.id
    
    # Get user's vector store
    vector_store = get_or_create_vector_store(user_id)
    
    if vector_store is None:
        raise HTTPException(
            status_code=400,
            detail="No documents uploaded yet. Please upload PDFs first."
        )
    
    openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
    if not openrouter_api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENROUTER_API_KEY not configured"
        )
    
    try:
        # Initialize LLM
        llm = ChatOpenAI(
            api_key=openrouter_api_key,
            base_url="https://openrouter.ai/api/v1",
            model="google/gemini-2.0-flash-001",
            temperature=0.1,
            max_tokens=8192,
        )
        
        # Get ALL documents for complete summary
        # First get collection to know total documents
        try:
            collection = vector_store.get()
            total_docs = len(collection.get('ids', []))
        except:
            # If get() fails, try alternative method
            total_docs = 100
        
        k_docs = min(total_docs, 100)  # Get up to 100 docs or all if less
        
        print(f"Generating summary for {total_docs} document chunks...")
        
        retriever = vector_store.as_retriever(search_kwargs={"k": k_docs})
        docs = retriever.invoke("summarize all content key points main topics")
        
        if not docs:
            raise HTTPException(
                status_code=400,
                detail="No content available to summarize"
            )
        
        # Join all content - take more for complete summary
        context = "\n\n".join(doc.page_content for doc in docs)
        
        # STRICT SYSTEM PROMPT for summary
        system_prompt = (
            "You are a strict data-extraction AI. Generate a comprehensive but concise summary of the provided document context. "
            "Focus on key points, main topics, and important details. "
            "Do NOT use your internal knowledge - ONLY use what is in the provided context.\n\n"
            "Context:\n"
            "{context}"
        )
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "Provide a detailed summary of this document."),
        ])
        
        # Create simple chain for summary
        summary_chain = prompt | llm | StrOutputParser()
        summary = summary_chain.invoke({"context": context})
        
        return SummaryResponse(summary=summary)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error during summary generation: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Convert Summary to PDF ---
@app.post("/summary/convert/pdf")
async def convert_summary_to_pdf(
    summary_text: str = Body(...),
    current_user = Depends(get_current_user)
):
    """Convert summary text to PDF and return as downloadable file."""
    try:
        from weasyprint import HTML
        import io
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; padding: 40px; line-height: 1.6; }}
                h1 {{ color: #333; }}
                p {{ text-align: justify; }}
            </style>
        </head>
        <body>
            <h1>Document Summary</h1>
            <p>{summary_text.replace(chr(10), '<br>')}</p>
        </body>
        </html>
        """
        
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=summary.pdf"}
        )
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Convert Summary to Word ---
@app.post("/summary/convert/word")
async def convert_summary_to_word(
    summary_text: str = Body(...),
    current_user = Depends(get_current_user)
):
    """Convert summary text to Word document and return as downloadable file."""
    try:
        from docx import Document
        import io
        
        doc = Document()
        doc.add_heading('Document Summary', 0)
        
        # Add paragraphs
        for para in summary_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=summary.docx"}
        )
    except Exception as e:
        print(f"Error converting to Word: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- User Documents Endpoints ---
@app.get("/documents")
async def get_user_documents(
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all documents uploaded by the current user."""
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()
    return {
        "user_id": current_user.id,
        "documents": [
            {
                "id": doc.id,
                "filename": doc.original_filename,
                "chunk_count": doc.chunk_count,
                "created_at": doc.created_at
            }
            for doc in documents
        ]
    }

@app.get("/history")
async def get_query_history(
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db),
    limit: int = 50
):
    """Get user's query history."""
    query_logs = db.query(QueryLog).filter(
        QueryLog.user_id == current_user.id
    ).order_by(QueryLog.created_at.desc()).limit(limit).all()
    
    return {
        "user_id": current_user.id,
        "history": [
            {
                "id": log.id,
                "document_id": log.document_id,
                "query": log.query,
                "response": log.response[:200] + "..." if len(log.response) > 200 else log.response,
                "created_at": log.created_at
            }
            for log in query_logs
        ]
    }

# --- Highlights Endpoints ---
class HighlightCreate(BaseModel):
    document_id: int
    text: str
    color: str = "#FFEB3B"
    page_number: int = 1

class HighlightResponse(BaseModel):
    id: int
    document_id: int
    text: str
    color: str
    page_number: int
    created_at: datetime

@app.post("/highlights", response_model=HighlightResponse)
async def create_highlight(
    highlight: HighlightCreate,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new highlight."""
    new_highlight = Highlight(
        user_id=current_user.id,
        document_id=highlight.document_id,
        text=highlight.text,
        color=highlight.color,
        page_number=highlight.page_number
    )
    db.add(new_highlight)
    db.commit()
    db.refresh(new_highlight)
    return new_highlight

@app.get("/highlights")
async def get_highlights(
    document_id: int = None,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all highlights for the user."""
    query = db.query(Highlight).filter(Highlight.user_id == current_user.id)
    if document_id:
        query = query.filter(Highlight.document_id == document_id)
    highlights = query.all()
    return {
        "highlights": [
            {
                "id": h.id,
                "document_id": h.document_id,
                "text": h.text,
                "color": h.color,
                "page_number": h.page_number,
                "created_at": h.created_at
            }
            for h in highlights
        ]
    }

@app.delete("/highlights/{highlight_id}")
async def delete_highlight(
    highlight_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a highlight."""
    highlight = db.query(Highlight).filter(
        Highlight.id == highlight_id,
        Highlight.user_id == current_user.id
    ).first()
    if not highlight:
        raise HTTPException(status_code=404, detail="Highlight not found")
    db.delete(highlight)
    db.commit()
    return {"message": "Highlight deleted"}

# --- Notes Endpoints ---
class NoteCreate(BaseModel):
    document_id: int
    content: str
    page_number: int = 1

class NoteResponse(BaseModel):
    id: int
    document_id: int
    content: str
    page_number: int
    created_at: datetime

@app.post("/notes", response_model=NoteResponse)
async def create_note(
    note: NoteCreate,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new note."""
    new_note = Note(
        user_id=current_user.id,
        document_id=note.document_id,
        content=note.content,
        page_number=note.page_number
    )
    db.add(new_note)
    db.commit()
    db.refresh(new_note)
    return new_note

@app.get("/notes")
async def get_notes(
    document_id: int = None,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all notes for the user."""
    query = db.query(Note).filter(Note.user_id == current_user.id)
    if document_id:
        query = query.filter(Note.document_id == document_id)
    notes = query.all()
    return {
        "notes": [
            {
                "id": n.id,
                "document_id": n.document_id,
                "content": n.content,
                "page_number": n.page_number,
                "created_at": n.created_at
            }
            for n in notes
        ]
    }

@app.delete("/notes/{note_id}")
async def delete_note(
    note_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a note."""
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    db.delete(note)
    db.commit()
    return {"message": "Note deleted"}

# --- Health Check ---
@app.get("/")
async def root():
    """Health check endpoint."""
    return {
        "message": "AI PDF Chatbot Backend v2.0 is running",
        "status": "online",
        "features": ["JWT Auth", "Multi-PDF Support", "Strict LLM Scoping"]
    }
